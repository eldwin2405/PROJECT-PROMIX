from pathlib import Path
import csv
import json
import re
from io import BytesIO

import streamlit as st
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

INCLUDE_MIE_DOT_SECTION = False
INCLUDE_M41 = False

REQUESTED_ORDER = [
    "Mie Suit",
    "Mie Gacoan Level 0",
    "Mie Gacoan Level 1",
    "Mie Gacoan Level 2",
    "Mie Gacoan Level 3",
    "Mie Gacoan Level 4",
    "Mie Gacoan Level 6",
    "Mie Gacoan Level 8",
    "Mie Hompimpa Level 1",
    "Mie Hompimpa Level 2",
    "Mie Hompimpa Level 3",
    "Mie Hompimpa Level 4",
    "Mie Hompimpa Level 6",
    "Mie Hompimpa Level 8",
    "Udang Keju",
    "Udang Rambutan",
    "Siomay",
    "Lumpia Udang",
    "Pangsit Goreng",
    "Air Mineral",
    "Lemon Tea Hot",
    "Lemon Tea Ice",
    "Total Lemon Tea",
    "Orange Hot",
    "Orange Ice",
    "Total Orange",
    "Teh Tarik Hot",
    "Teh Tarik Ice",
    "Total Teh Tarik",
    "Milo Hot",
    "Milo Ice",
    "Total Milo",
    "Vanilla Latte Hot",
    "Vanilla Latte Ice",
    "Total Vanilla Latte",
    "Tea Hot",
    "Tea Ice",
    "Total Tea",
    "Es Gobak Sodor",
    "Es Teklek",
    "Es Petak Umpet",
    "Es Sluku Bathok",
    "Thai Tea",
    "Thai Green Tea",
    "Lemon Tea Ice NP",
    "Orange Ice NP",
    "Teh Tarik Ice NP",
    "Vanilla Latte Ice NP",
    "Lemon Tea Hot NP",
    "Orange Hot NP",
    "Teh Tarik Hot NP",
    "Vanilla Latte Hot NP",
    "Cup 16",
    "Packaging Mie",
    "Packaging Dimsum",
    "Total Packaging",
]

CATEGORY_GROUPS = {
    "Mie": [
        "Mie Suit",
        "Mie Gacoan Level 0",
        "Mie Gacoan Level 1",
        "Mie Gacoan Level 2",
        "Mie Gacoan Level 3",
        "Mie Gacoan Level 4",
        "Mie Gacoan Level 6",
        "Mie Gacoan Level 8",
        "Mie Hompimpa Level 1",
        "Mie Hompimpa Level 2",
        "Mie Hompimpa Level 3",
        "Mie Hompimpa Level 4",
        "Mie Hompimpa Level 6",
        "Mie Hompimpa Level 8",
    ],
    "Dimsum": [
        "Udang Keju",
        "Udang Rambutan",
        "Siomay",
        "Lumpia Udang",
        "Pangsit Goreng",
    ],
    "Beverages": [
        "Air Mineral",
        "Lemon Tea Hot",
        "Lemon Tea Ice",
        "Total Lemon Tea",
        "Orange Hot",
        "Orange Ice",
        "Total Orange",
        "Teh Tarik Hot",
        "Teh Tarik Ice",
        "Total Teh Tarik",
        "Milo Hot",
        "Milo Ice",
        "Total Milo",
        "Vanilla Latte Hot",
        "Vanilla Latte Ice",
        "Total Vanilla Latte",
        "Tea Hot",
        "Tea Ice",
        "Total Tea",
        "Thai Tea",
        "Thai Green Tea",
    ],
    "Es Buah": [
        "Es Gobak Sodor",
        "Es Teklek",
        "Es Petak Umpet",
        "Es Sluku Bathok",
    ],
    "NP": [
        "Lemon Tea Ice NP",
        "Orange Ice NP",
        "Teh Tarik Ice NP",
        "Vanilla Latte Ice NP",
        "Lemon Tea Hot NP",
        "Orange Hot NP",
        "Teh Tarik Hot NP",
        "Vanilla Latte Hot NP",
    ],
    "Packaging": [
        "Cup 16",
        "Packaging Mie",
        "Packaging Dimsum",
        "Total Packaging",
    ],
}

TOP_LEVEL_BLOCKS = {
    "BEVERAGES",
    "BEVERAGES ICE.",
    "BEVERAGES ISIAN",
    "DIMSUM",
    "ES BUAH",
    "MIE",
    "MIE ISIAN",
    "MIE.",
    "PACKAGING",
    "PAKET",
}


def extract_all_docx_text(file_obj) -> str:
    doc = Document(file_obj)
    lines = []

    lines.append("=== PARAGRAPHS ===")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    lines.append("\n=== TABLES ===")
    for table_index, table in enumerate(doc.tables, start=1):
        lines.append(f"\n--- TABLE {table_index} ---")
        for row in table.rows:
            row_values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(" | ".join(row_values))

    return "\n".join(lines)


def make_empty_result():
    return {
        name: {"menu": name, "dine_in": 0, "take_away": 0, "total": 0}
        for name in REQUESTED_ORDER
    }


def clean_qty(text: str) -> int:
    text = text.strip()
    if not text:
        return 0
    return int(text.replace(".", ""))


def add_qty(results: dict, menu: str, channel: str, qty: int) -> None:
    if menu not in results:
        return
    results[menu][channel] += qty


def extract_table_block(text: str, table_number: int) -> list[str]:
    pattern = rf"--- TABLE {table_number} ---\n(.*?)(?=\n--- TABLE \d+ ---|\Z)"
    match = re.search(pattern, text, flags=re.S)
    if not match:
        raise ValueError(f"TABLE {table_number} tidak ditemukan di hasil ekstraksi.")

    lines = []
    for line in match.group(1).splitlines():
        line = line.strip()
        if line:
            lines.append(line)
    return lines


def find_sales_by_menu_table_number(text: str) -> int:
    for table_number in range(1, 30):
        try:
            lines = extract_table_block(text, table_number)
        except ValueError:
            continue

        block_text = "\n".join(lines)

        if (
            "Description | Qty | Value" in block_text
            and "BEVERAGES" in block_text
            and "MIE ISIAN" in block_text
            and "PACKAGING" in block_text
        ):
            return table_number

    raise ValueError("Tabel Sales By Menu tidak ditemukan.")


def compute_totals(results: dict) -> None:
    for row in results.values():
        row["total"] = row["dine_in"] + row["take_away"]

    grouped_totals = [
        ("Total Lemon Tea", "Lemon Tea Hot", "Lemon Tea Ice"),
        ("Total Orange", "Orange Hot", "Orange Ice"),
        ("Total Teh Tarik", "Teh Tarik Hot", "Teh Tarik Ice"),
        ("Total Milo", "Milo Hot", "Milo Ice"),
        ("Total Vanilla Latte", "Vanilla Latte Hot", "Vanilla Latte Ice"),
        ("Total Tea", "Tea Hot", "Tea Ice"),
    ]

    for total_name, hot_name, ice_name in grouped_totals:
        results[total_name]["dine_in"] = (
            results[hot_name]["dine_in"] + results[ice_name]["dine_in"]
        )
        results[total_name]["take_away"] = (
            results[hot_name]["take_away"] + results[ice_name]["take_away"]
        )
        results[total_name]["total"] = (
            results[total_name]["dine_in"] + results[total_name]["take_away"]
        )

    if results["Total Packaging"]["total"] == 0:
        results["Total Packaging"]["dine_in"] = (
            results["Cup 16"]["dine_in"]
            + results["Packaging Mie"]["dine_in"]
            + results["Packaging Dimsum"]["dine_in"]
        )
        results["Total Packaging"]["take_away"] = (
            results["Cup 16"]["take_away"]
            + results["Packaging Mie"]["take_away"]
            + results["Packaging Dimsum"]["take_away"]
        )
        results["Total Packaging"]["total"] = (
            results["Total Packaging"]["dine_in"]
            + results["Total Packaging"]["take_away"]
        )


def parse_target_data(lines: list[str]) -> dict:
    results = make_empty_result()

    current_block = None
    current_subblock = None

    for line in lines[1:]:
        parts = [p.strip() for p in line.split("|")]
        if len(parts) != 3:
            continue

        description, qty_text, value_text = parts

        is_header = qty_text == "" and value_text == ""
        if is_header:
            if description in TOP_LEVEL_BLOCKS:
                current_block = description
                current_subblock = None
            else:
                current_subblock = description
            continue

        qty = clean_qty(qty_text)

        if current_block == "PACKAGING" and description == "Total - PACKAGING":
            add_qty(results, "Total Packaging", "dine_in", qty)
            continue

        if description.startswith("Total -"):
            continue

        if current_block == "BEVERAGES":
            if current_subblock == "BEVERAGES DINE IN":
                channel = "dine_in"
            elif current_subblock == "BEVERAGES TAKE AWAY":
                channel = "take_away"
            else:
                channel = None

            direct_beverage_map = {
                "AIR MINERAL": "Air Mineral",
                "THAI TEA": "Thai Tea",
                "THAI GREEN TEA": "Thai Green Tea",
            }

            if channel and description in direct_beverage_map:
                add_qty(results, direct_beverage_map[description], channel, qty)

        elif current_block == "BEVERAGES ISIAN":
            beverage_isian_map = {
                "LEMON TEA DINE IN": ("dine_in", {"HOT": "Lemon Tea Hot", "ICED": "Lemon Tea Ice"}),
                "LEMON TEA TA": ("take_away", {"HOT": "Lemon Tea Hot", "ICED": "Lemon Tea Ice"}),
                "ORANGE DINE IN": ("dine_in", {"HOT": "Orange Hot", "ICED": "Orange Ice"}),
                "ORANGE TA": ("take_away", {"HOT": "Orange Hot", "ICED": "Orange Ice"}),
                "TEH TARIK DINE IN": ("dine_in", {"HOT": "Teh Tarik Hot", "ICED": "Teh Tarik Ice"}),
                "TEH TARIK TA": ("take_away", {"HOT": "Teh Tarik Hot", "ICED": "Teh Tarik Ice"}),
                "MILO DINE IN": ("dine_in", {"HOT": "Milo Hot", "ICED": "Milo Ice"}),
                "MILO TA": ("take_away", {"HOT": "Milo Hot", "ICED": "Milo Ice"}),
                "VANILLA LATTE DINE IN": ("dine_in", {"HOT": "Vanilla Latte Hot", "ICED": "Vanilla Latte Ice"}),
                "VANILLA LATTE TA": ("take_away", {"HOT": "Vanilla Latte Hot", "ICED": "Vanilla Latte Ice"}),
                "TEA DINE IN": ("dine_in", {"HOT": "Tea Hot", "ICED": "Tea Ice"}),
                "TEA TA": ("take_away", {"HOT": "Tea Hot", "ICED": "Tea Ice"}),
                "VANILLA LATTE NP DINE IN": (
                    "dine_in",
                    {"HOT": "Vanilla Latte Hot NP", "ICED": "Vanilla Latte Ice NP"},
                ),
                "VANILLA LATTE NP TA": (
                    "take_away",
                    {"HOT": "Vanilla Latte Hot NP", "ICED": "Vanilla Latte Ice NP"},
                ),
                "LEMON TEA NP DINE IN": (
                    "dine_in",
                    {"HOT": "Lemon Tea Hot NP", "ICED": "Lemon Tea Ice NP"},
                ),
                "LEMON TEA NP TA": (
                    "take_away",
                    {"HOT": "Lemon Tea Hot NP", "ICED": "Lemon Tea Ice NP"},
                ),
                "ORANGE NP DINE IN": (
                    "dine_in",
                    {"HOT": "Orange Hot NP", "ICED": "Orange Ice NP"},
                ),
                "ORANGE NP TA": (
                    "take_away",
                    {"HOT": "Orange Hot NP", "ICED": "Orange Ice NP"},
                ),
                "TEH TARIK NP DINE IN": (
                    "dine_in",
                    {"HOT": "Teh Tarik Hot NP", "ICED": "Teh Tarik Ice NP"},
                ),
                "TEH TARIK NP TA": (
                    "take_away",
                    {"HOT": "Teh Tarik Hot NP", "ICED": "Teh Tarik Ice NP"},
                ),
            }

            if current_subblock in beverage_isian_map:
                channel, item_map = beverage_isian_map[current_subblock]
                if description in item_map:
                    add_qty(results, item_map[description], channel, qty)

        elif current_block == "DIMSUM":
            if current_subblock == "DIMSUM DINE IN":
                channel = "dine_in"
            elif current_subblock == "DIMSUM TAKE AWAY":
                channel = "take_away"
            else:
                channel = None

            dimsum_map = {
                "UDANG KEJU": "Udang Keju",
                "UDANG RAMBUTAN": "Udang Rambutan",
                "SIOMAY": "Siomay",
                "LUMPIA UDANG": "Lumpia Udang",
                "PANGSIT GORENG": "Pangsit Goreng",
            }

            if channel and description in dimsum_map:
                add_qty(results, dimsum_map[description], channel, qty)

        elif current_block == "ES BUAH":
            if current_subblock == "ES BUAH DINE IN":
                channel = "dine_in"
            elif current_subblock == "ES BUAH TAKE AWAY":
                channel = "take_away"
            else:
                channel = None

            es_map = {
                "ES GOBAK SODOR": "Es Gobak Sodor",
                "ES TEKLEK": "Es Teklek",
                "ES PETAK UMPET": "Es Petak Umpet",
                "ES SLUKU BATHOK": "Es Sluku Bathok",
            }

            if channel and description in es_map:
                add_qty(results, es_map[description], channel, qty)

        elif current_block == "MIE":
            if current_subblock == "MIE DINE IN":
                channel = "dine_in"
            elif current_subblock == "MIE TAKE AWAY":
                channel = "take_away"
            else:
                channel = None

            if channel and description == "MIE SUIT":
                add_qty(results, "Mie Suit", channel, qty)

        elif current_block == "MIE ISIAN":
            mie_isian_map = {
                "MIE GACOAN DINE IN": (
                    "dine_in",
                    {
                        "LEVEL 0": "Mie Gacoan Level 0",
                        "LEVEL 1": "Mie Gacoan Level 1",
                        "LEVEL 2": "Mie Gacoan Level 2",
                        "LEVEL 3": "Mie Gacoan Level 3",
                        "LEVEL 4": "Mie Gacoan Level 4",
                        "LEVEL 6": "Mie Gacoan Level 6",
                        "LEVEL 8": "Mie Gacoan Level 8",
                    },
                ),
                "MIE GACOAN TAKE AWAY": (
                    "take_away",
                    {
                        "LEVEL 0": "Mie Gacoan Level 0",
                        "LEVEL 1": "Mie Gacoan Level 1",
                        "LEVEL 2": "Mie Gacoan Level 2",
                        "LEVEL 3": "Mie Gacoan Level 3",
                        "LEVEL 4": "Mie Gacoan Level 4",
                        "LEVEL 6": "Mie Gacoan Level 6",
                        "LEVEL 8": "Mie Gacoan Level 8",
                    },
                ),
                "MIE HOMPIMPA DINE IN": (
                    "dine_in",
                    {
                        "LEVEL 1": "Mie Hompimpa Level 1",
                        "LEVEL 2": "Mie Hompimpa Level 2",
                        "LEVEL 3": "Mie Hompimpa Level 3",
                        "LEVEL 4": "Mie Hompimpa Level 4",
                        "LEVEL 6": "Mie Hompimpa Level 6",
                        "LEVEL 8": "Mie Hompimpa Level 8",
                    },
                ),
                "MIE HOMPIMPA TAKE AWAY": (
                    "take_away",
                    {
                        "LEVEL 1": "Mie Hompimpa Level 1",
                        "LEVEL 2": "Mie Hompimpa Level 2",
                        "LEVEL 3": "Mie Hompimpa Level 3",
                        "LEVEL 4": "Mie Hompimpa Level 4",
                        "LEVEL 6": "Mie Hompimpa Level 6",
                        "LEVEL 8": "Mie Hompimpa Level 8",
                    },
                ),
            }

            if current_subblock in mie_isian_map:
                channel, item_map = mie_isian_map[current_subblock]
                if description in item_map:
                    add_qty(results, item_map[description], channel, qty)

        elif current_block == "MIE." and INCLUDE_MIE_DOT_SECTION:
            direct_mie_dot_map = {
                "MIE SUIT": "Mie Suit",
                "MIE GACOAN LEVEL 0": "Mie Gacoan Level 0",
                "MIE GACOAN LEVEL 1": "Mie Gacoan Level 1",
                "MIE HOMPIMPA LEVEL 1": "Mie Hompimpa Level 1",
            }

            if current_subblock == "MIE DINE IN":
                if description in direct_mie_dot_map:
                    add_qty(results, direct_mie_dot_map[description], "dine_in", qty)
            elif current_subblock == "MIE TAKE AWAY":
                if description in direct_mie_dot_map:
                    add_qty(results, direct_mie_dot_map[description], "take_away", qty)
            elif current_subblock == "MIE GRAB FOOD" and INCLUDE_M41:
                m41_map = {
                    "MIE GACOAN LEVEL 0 - M41": "Mie Gacoan Level 0",
                    "MIE GACOAN LEVEL 1 - M41": "Mie Gacoan Level 1",
                    "MIE GACOAN LEVEL 3 - M41": "Mie Gacoan Level 3",
                }
                if description in m41_map:
                    add_qty(results, m41_map[description], "take_away", qty)

        elif current_block == "PACKAGING":
            if current_subblock == "PACKAGING DINE IN":
                packaging_map = {
                    "16 OZ": "Cup 16",
                    "MIE": "Packaging Mie",
                    "DIMSUM": "Packaging Dimsum",
                }
                if description in packaging_map:
                    add_qty(results, packaging_map[description], "dine_in", qty)

    compute_totals(results)
    return results


def grouped_json_data(results: dict):
    grouped_data = []
    for category, items in CATEGORY_GROUPS.items():
        grouped_data.append({
            "category": category,
            "items": [results[name] for name in items]
        })
    return grouped_data


def csv_bytes(results: dict) -> bytes:
    bio = BytesIO()
    text = []
    writer = csv.writer(text := [])
    # csv.writer needs file-like object; use StringIO instead
    raise RuntimeError("unused")


def make_csv_text(results: dict) -> str:
    from io import StringIO
    buffer = StringIO()
    writer = csv.writer(buffer)
    for category, items in CATEGORY_GROUPS.items():
        writer.writerow([category])
        writer.writerow(["Menu", "Dine In", "Take Away", "Total"])
        for name in items:
            row = results[name]
            menu_name = row["menu"].upper() if row["menu"].startswith("Total") else row["menu"]
            writer.writerow([menu_name, row["dine_in"], row["take_away"], row["total"]])
        writer.writerow([])
    return buffer.getvalue()


def build_category_table(results: dict, items: list[str]) -> Table:
    table_data = [["Menu", "Dine In", "Take Away", "Total"]]
    total_row_indexes = []

    for name in items:
        row = results[name]
        menu_name = row["menu"]
        if menu_name.startswith("Total"):
            menu_name = menu_name.upper()
            total_row_indexes.append(len(table_data))

        table_data.append([
            menu_name,
            str(row["dine_in"]),
            str(row["take_away"]),
            str(row["total"]),
        ])

    table = Table(table_data, repeatRows=1, colWidths=[220, 70, 70, 70])
    style_commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]
    for row_index in total_row_indexes:
        style_commands.append(("FONTNAME", (0, row_index), (-1, row_index), "Helvetica-Bold"))
    table.setStyle(TableStyle(style_commands))
    return table


def make_pdf_bytes(results: dict, source_name: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Laporan Ekstraksi Menu PROMIX", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Sumber file: {source_name}", styles["Normal"]))
    story.append(Spacer(1, 16))

    categories = list(CATEGORY_GROUPS.items())
    for index, (category, items) in enumerate(categories):
        story.append(Paragraph(category, styles["Heading2"]))
        story.append(Spacer(1, 6))
        story.append(build_category_table(results, items))
        if index < len(categories) - 1:
            story.append(PageBreak())

    doc.build(story)
    return buffer.getvalue()


def grouped_frames(results: dict):
    import pandas as pd
    frames = {}
    for category, items in CATEGORY_GROUPS.items():
        rows = []
        for name in items:
            row = results[name]
            rows.append({
                "Menu": row["menu"].upper() if row["menu"].startswith("Total") else row["menu"],
                "Dine In": row["dine_in"],
                "Take Away": row["take_away"],
                "Total": row["total"],
            })
        frames[category] = pd.DataFrame(rows)
    return frames


def process_docx(file_obj, file_name: str):
    extracted_text = extract_all_docx_text(file_obj)
    table_number = find_sales_by_menu_table_number(extracted_text)
    table_lines = extract_table_block(extracted_text, table_number)
    results = parse_target_data(table_lines)
    return extracted_text, table_number, results


st.set_page_config(page_title="PROMIX Streamlit Extractor", layout="wide")
st.title("CICILAN NGEBUT PROMIX")
st.caption("Upload 1 file DOCX laporan PROMIX")

with st.sidebar:
    st.subheader("Pengaturan")
    st.write(f"INCLUDE_MIE_DOT_SECTION: {INCLUDE_MIE_DOT_SECTION}")
    st.write(f"INCLUDE_M41: {INCLUDE_M41}")
    st.info("Versi ini membaca tabel Sales By Menu secara otomatis tanpa bergantung pada nomor tabel tetap.")

uploaded_file = st.file_uploader("Upload file DOCX", type=["docx"])

if uploaded_file is not None:
    try:
        extracted_text, table_number, results = process_docx(uploaded_file, uploaded_file.name)
        
        st.success(f"Berhasil diproses. Sales By Menu ditemukan di TABLE {table_number}.")

        st.divider()
        st.subheader("Preview Hasil per Kategori")
        frames = grouped_frames(results)
        for category, df in frames.items():
            with st.expander(category, expanded=(category == "Mie")):
                st.dataframe(df, use_container_width=True, hide_index=True)

    except Exception as e:
        st.error(f"Terjadi error: {e}")
        st.exception(e)
else:
    st.info("Silakan upload 1 file DOCX terlebih dahulu.")
